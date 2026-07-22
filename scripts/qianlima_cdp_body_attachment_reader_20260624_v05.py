# -*- coding: utf-8 -*-
"""Qianlima CDP body and attachment reader.

Connect to an already logged-in Chrome through Playwright CDP. Read the real
project body, discover attachments, resolve Qianlima's two-step download URLs,
parse supported files, and produce one output record for every input project.

This reader is conservative: short/scanned PDFs, unsupported archives, missing
attachments, login pages, and identity mismatches are not successful evidence.
"""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import re
import subprocess
import sys
import time
import traceback
import zipfile
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import urljoin, urlparse

import requests

try:
    from playwright.sync_api import TimeoutError as PlaywrightTimeoutError
    from playwright.sync_api import sync_playwright
except ImportError as exc:
    raise SystemExit("Missing dependency: playwright. Run: pip install playwright") from exc

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None


MIN_BODY_CHARS = 200
MIN_ATTACHMENT_CHARS = 80
MAX_FILE_BYTES = 200 * 1024 * 1024  # 2026-07-10 提高到 200MB
# 附件文本类（用于 ZIP 内/项目级「未读」判定）；图纸/图片不挡整包成功
TEXT_BEARING_TYPES = frozenset({"docx", "doc", "pdf", "xlsx", "xls", "txt", "csv", "xml", "html", "htm", "rtf"})
# 名称含下列关键词且 OCR/解析失败时，视为非主证据（签章/委托），不单独把 attachment_read_ok 打成失败
OPTIONAL_ATTACH_NAME_HINTS = (
    "声明", "委托书", "委托", "签章", "签名", "电子签名", "营业执照", "身份证", "授权书",
    "图纸", "CAD", "cad", "效果图", "施工图", "蓝图",
)
# ZIP 内跳过、不计入 unresolved 的类型/扩展
ZIP_SKIP_NAME_EXT = (
    ".dwg", ".dxf", ".dwf", ".rvt", ".skp", ".png", ".jpg", ".jpeg", ".gif",
    ".bmp", ".tif", ".tiff", ".webp", ".mp4", ".avi", ".exe", ".dll",
)
# 2026-07-14：原 10MB 跳过 OCR 过严（大附件直接不读字）。
# 提高到与下载上限同量级；仅极端大文件跳过整段 OCR。
LARGE_FILE_SKIP_OCR = 100 * 1024 * 1024  # 超过 100MB 跳过 OCR（原 10MB）
# OCR 进度心跳：无进展超过 BASE 则停；有进展可续到 HARD_CAP
OCR_POLL_INTERVAL_SEC = 30
OCR_STALL_TIMEOUT_SEC = 300   # base：连续无页进度则终止（默认 300s）
OCR_HARD_CAP_SEC = 1800       # hard cap：单附件 OCR 最长 30 分钟
# 2026-07-14 P0（标神工单 attach-download-watchdog）：分段日志 + 下载/预览看门狗 + 预算
DOWNLOAD_CONNECT_TIMEOUT_SEC = 30
DOWNLOAD_READ_TIMEOUT_SEC = 90          # 原 requests 读超时 600s 过长
DOWNLOAD_STALL_SEC = 60                 # 字节无增长则判死
DOWNLOAD_HARD_CAP_SEC = 300             # 单次流式下载硬顶 5 分钟
PREVIEW_POLL_MAX_SEC = 60               # PDF.js 等待硬顶（原最多 90s 死循环）
PREVIEW_STALL_SEC = 45                  # 预览无 ready 且无字节/状态变化
SPA_POLL_MAX_SEC = 45                   # SPA 回退轮询硬顶（原 30×1s）
PER_ATTACHMENT_BUDGET_SEC = 180         # 单附件总预算（下载+解析+OCR 分段另计）
PER_PROJECT_BUDGET_SEC = 420            # 单项目全部附件总预算
MAX_ZIP_FILES = 60
MAX_ZIP_UNCOMPRESSED_BYTES = 150 * 1024 * 1024

_SEG_LOG: Optional["SegmentLogger"] = None


class SegmentLogger:
    """Append-only phase2 attach log → runs/<id>/99_logs/ when layout matches."""

    def __init__(self, path: Path):
        self.path = path
        self.path.parent.mkdir(parents=True, exist_ok=True)
        self._fp = self.path.open("a", encoding="utf-8", errors="replace")
        self.line(f"=== segment log open {time.strftime('%Y-%m-%d %H:%M:%S')} path={path}")

    def line(self, msg: str) -> None:
        ts = time.strftime("%H:%M:%S")
        text = f"[{ts}] {msg}"
        print(text, flush=True)
        try:
            self._fp.write(text + "\n")
            self._fp.flush()
        except Exception:
            pass

    def close(self) -> None:
        try:
            self.line("=== segment log close ===")
            self._fp.close()
        except Exception:
            pass


def slog(msg: str) -> None:
    """stdout + optional segment file."""
    if _SEG_LOG is not None:
        _SEG_LOG.line(msg)
    else:
        print(msg, flush=True)


def resolve_phase2_log_dir(output_dir: Path) -> Path:
    """Prefer runs/<run_id>/99_logs when output is .../03_body/cdp_evidence."""
    p = output_dir.resolve()
    if p.name == "cdp_evidence" and p.parent.name == "03_body":
        return p.parent.parent / "99_logs"
    return p / "99_logs"


def budget_exceeded(start: float, budget_sec: float) -> bool:
    return (time.time() - start) >= budget_sec
LOGIN_TEXT = ("登录状态超时", "请重新登录", "登录状态已过期")
CAPTCHA_TEXT = ("Access Verification", "安全验证", "验证码", "访问验证")
BODY_STOP_MARKERS = ("企业商情分析", "招标进度跟踪", "相关招标", "热门推荐", "产品意见征集")
ATTACHMENT_HINTS = ("详见附件", "以附件为准", "采购文件", "招标文件", "技术规范书")
ATTACHMENT_NAME_HINTS = (
    "附件", "采购文件", "招标文件", "询价文件", "竞价文件", "需求书",
    "技术规范", "技术条件", "工程量清单", "图纸", "合同", "询价函",
)
NON_EVIDENCE_ACTION_NAMES = (
    "导出Word", "导出Excel", "导出PDF", "供应商报名", "投标人登记申请表",
    "企业信息", "立即查看", "立即荐标", "立即监控",
)
SUPPORTED_DOCUMENT_SUFFIXES = (
    ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".zip", ".rar", ".7z",
)
STATIC_RESOURCE_SUFFIXES = (
    ".css", ".js", ".mjs", ".map", ".woff", ".woff2", ".ttf", ".ico",
)


def compact_text(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or ""))


def normalize_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def clean_attachment_display_name(value: str) -> str:
    """去掉「（86KB）/ (9KB)」等体积后缀，便于识别扩展名。"""
    name = normalize_text(value)
    name = re.sub(r"[\(（]\s*\d+(?:\.\d+)?\s*[KMG]?B\s*[\)）]\s*$", "", name, flags=re.I)
    name = re.sub(r"^附件[-：:\s]+", "", name)
    return name.strip(" .") or normalize_text(value)


def safe_filename(value: str, fallback: str = "attachment") -> str:
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", clean_attachment_display_name(value))
    return (name.strip(" .")[:160] or fallback)


def has_meaningful_text(value: Any, minimum: int) -> bool:
    compact = compact_text(value)
    return len(compact) >= minimum and len(set(compact)) >= 8


def extract_bid_id(project: Dict[str, Any]) -> str:
    direct = normalize_text(project.get("bid_id"))
    if direct:
        return direct
    for key in ("url", "detail_url", "link", "千里马链接"):
        match = re.search(r"bid-(\d+)", str(project.get(key, "")), re.I)
        if match:
            return f"bid-{match.group(1)}"
    return ""


def project_url(project: Dict[str, Any]) -> str:
    for key in ("url", "detail_url", "link", "千里马链接"):
        value = normalize_text(project.get(key))
        if value.startswith(("http://", "https://")):
            return value
    bid_id = extract_bid_id(project)
    return f"https://www.qianlima.com/{bid_id}.html" if bid_id else ""


def infer_file_type(filename: str, content_type: str = "") -> str:
    cleaned = clean_attachment_display_name(filename or "")
    # 名称里可能带路径或 query，优先取 basename 扩展名
    base = Path(cleaned.replace("\\", "/").split("/")[-1]).name
    ext = Path(base).suffix.lower()
    if not ext:
        ext = Path(urlparse(cleaned).path).suffix.lower()
    mapping = {
        ".docx": "docx", ".doc": "doc", ".pdf": "pdf",
        ".xlsx": "xlsx", ".xls": "xls", ".zip": "zip",
        ".rar": "rar", ".7z": "7z", ".png": "image",
        ".jpg": "image", ".jpeg": "image", ".webp": "image",
        ".tif": "image", ".tiff": "image",
    }
    if ext in mapping:
        return mapping[ext]
    content_type = (content_type or "").lower()
    if "pdf" in content_type:
        return "pdf"
    if "wordprocessingml" in content_type or "msword" in content_type:
        return "docx" if "openxml" in content_type or "wordprocessingml" in content_type else "doc"
    if "spreadsheetml" in content_type or "excel" in content_type:
        return "xlsx" if "openxml" in content_type or "spreadsheetml" in content_type else "xls"
    if "zip" in content_type or "octet-stream" in content_type:
        # octet-stream 仍可能是任意二进制，留给 sniff
        if "zip" in content_type:
            return "zip"
    if content_type.startswith("image/"):
        return "image"
    return "unknown"


def sniff_file_type(content: bytes) -> str:
    """用魔数识别类型；名称/content-type 缺失时的 U2 回退。"""
    if not content:
        return "unknown"
    head = content[:16]
    if head.startswith(b"%PDF"):
        return "pdf"
    if head.startswith(b"PK"):
        # docx/xlsx 也是 zip 容器
        try:
            import io
            with zipfile.ZipFile(io.BytesIO(content[: min(len(content), 2 * 1024 * 1024)])) as zf:
                names = " ".join(zf.namelist()[:80]).lower()
                if "word/" in names:
                    return "docx"
                if "xl/" in names:
                    return "xlsx"
                if "ppt/" in names:
                    return "pptx"
        except Exception:
            pass
        return "zip"
    if head.startswith(bytes.fromhex("D0CF11E0A1B11AE1")):
        # OLE：doc/xls 粗分为 doc（解析层再处理）
        return "doc"
    if head[:3] == b"\xff\xd8\xff":
        return "image"
    if head.startswith(b"\x89PNG"):
        return "image"
    if head.startswith(b"Rar!"):
        return "rar"
    if head.startswith(b"7z\xbc\xaf\x27\x1c"):
        return "7z"
    return "unknown"


def identity_compact(value: Any) -> str:
    text = compact_text(value)
    text = re.sub(r"(企业信息|加急标书代写|标书代写)", "", text)
    return re.sub(r"[^\u4e00-\u9fffA-Za-z0-9]+", "", text)


def title_identity_candidates(title: str) -> List[str]:
    raw = normalize_text(title)
    variants = [raw]
    variants.extend(re.findall(r"【([^】]{6,})】", raw))
    stripped = re.sub(r"^(?:【[^】]+】)+", "", raw)
    stripped = re.sub(r"^(?:关于为|关于|采购项目名称[:：]?)", "", stripped)
    stripped = re.sub(
        r"(?:招标公告|采购公告|询比公告|竞价公告|公开选取.*?机构的公告)\s*$",
        "",
        stripped,
    )
    variants.append(stripped)
    candidates: List[str] = []
    for value in variants:
        candidate = identity_compact(value)
        if len(candidate) >= 8 and candidate not in candidates:
            candidates.append(candidate)
    return sorted(candidates, key=len, reverse=True)


def body_matches_project(title: str, body_text: str) -> bool:
    actual = identity_compact(body_text)
    candidates = title_identity_candidates(title)
    if not candidates or not actual:
        return False
    for candidate in candidates:
        if candidate in actual:
            return True
        match = SequenceMatcher(None, candidate[:120], actual).find_longest_match()
        threshold = max(8, min(16, int(len(candidate) * 0.55)))
        if match.size >= threshold:
            return True
    return False

def iter_projects(data: Any) -> List[Dict[str, Any]]:
    if isinstance(data, list):
        candidates = data
    elif isinstance(data, dict):
        candidates = []
        found = False
        for key in ("recommended", "considered", "low_score", "projects", "items", "results"):
            value = data.get(key)
            if isinstance(value, list):
                found = True
                candidates.extend(value)
        if not found and all(isinstance(value, dict) for value in data.values()):
            candidates = list(data.values())
    else:
        candidates = []
    output, seen = [], set()
    for item in candidates:
        if not isinstance(item, dict):
            continue
        key = extract_bid_id(item) or project_url(item) or normalize_text(item.get("title"))
        if key and key in seen:
            continue
        if key:
            seen.add(key)
        output.append(item)
    return output


@dataclass
class CDPClient:
    cdp_url: str
    screenshot_dir: Path

    def __post_init__(self) -> None:
        self.playwright = None
        self.browser = None
        self.context = None
        self.page = None

    def connect(self) -> None:
        self.playwright = sync_playwright().start()
        self.browser = self.playwright.chromium.connect_over_cdp(self.cdp_url)
        if not self.browser.contexts:
            raise RuntimeError("CDP Chrome has no browser context.")
        self.context = self.browser.contexts[0]
        self.page = self.context.new_page()
        self.page.set_default_timeout(20_000)

    def close(self) -> None:
        try:
            if self.page and not self.page.is_closed():
                self.page.close()
        finally:
            if self.playwright:
                self.playwright.stop()

    def navigate(self, url: str) -> None:
        last_error: Optional[Exception] = None
        for attempt in range(3):
            try:
                self.page.goto(url, wait_until="domcontentloaded", timeout=45_000)
                self.page.wait_for_timeout(1500)
                return
            except Exception as exc:
                last_error = exc
                if attempt < 2:
                    self.page.wait_for_timeout(1500 * (attempt + 1))
        raise RuntimeError(f"Navigation failed after 3 attempts: {last_error}")

    def visible_text(self) -> str:
        return self.page.locator("body").inner_text(timeout=10_000)

    def screenshot(self, name: str) -> str:
        path = self.screenshot_dir / f"{safe_filename(name)}.png"
        try:
            self.page.screenshot(path=str(path), full_page=True, timeout=15_000)
            return str(path)
        except Exception:
            return ""

    def cookies(self) -> Dict[str, str]:
        return {item["name"]: item["value"] for item in self.context.cookies()}


def classify_page_gate(text: str, url: str,
                       has_visible_password: bool = False) -> Tuple[str, str]:
    if any(keyword in text for keyword in CAPTCHA_TEXT):
        return "captcha", "页面出现验证码或访问验证"
    if "/login" in (url or "").lower():
        return "need_login", "页面位于登录地址"
    if any(keyword in text for keyword in LOGIN_TEXT):
        return "need_login", "页面出现明确登录失效提示"
    if has_visible_password:
        return "need_login", "页面显示密码登录表单"
    return "ok", ""


def page_gate(client: CDPClient) -> Tuple[str, str]:
    text = client.visible_text()
    url = client.page.url
    has_visible_password = False
    try:
        has_visible_password = (
            client.page.locator('input[type="password"]:visible').count() > 0
        )
    except Exception:
        pass
    return classify_page_gate(text, url, has_visible_password)

def dismiss_login_timeout_dialog(client: CDPClient) -> bool:
    """Close the known login-timeout dialog without touching browser data."""
    for selector in ("button", ".el-button", ".layui-layer-btn a", "span"):
        try:
            locator = client.page.locator(selector)
            for index in range(min(locator.count(), 30)):
                item = locator.nth(index)
                text = re.sub(r"\s+", "", item.inner_text(timeout=1_000))
                if text in {"确定", "关闭", "忽略"} and item.is_visible():
                    item.click(timeout=3_000)
                    client.page.wait_for_timeout(800)
                    return True
        except Exception:
            continue
    return False


def cdp_port_from_url(cdp_url: str) -> int:
    parsed = urlparse(cdp_url)
    return parsed.port or 9222


def restore_login_once(client: CDPClient, login_script: Path, config_path: Path | None) -> bool:
    """Restore the CDP account once, then let the caller retry the same project."""
    dismiss_login_timeout_dialog(client)
    try:
        client.navigate("https://vip.qianlima.com/")
        gate, _ = page_gate(client)
        if gate == "ok":
            return True
    except Exception:
        pass

    if not login_script.exists() or not config_path or not config_path.exists():
        return False
    cmd = [
        sys.executable,
        "-X",
        "utf8",
        str(login_script),
        "--mode",
        "cdp",
        "--cdp-port",
        str(cdp_port_from_url(client.cdp_url)),
        "--config",
        str(config_path),
    ]
    proc = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=120,
    )
    if proc.stdout:
        print(proc.stdout, flush=True)
    if proc.returncode != 0:
        if proc.stderr:
            print(proc.stderr, file=sys.stderr, flush=True)
        return False
    client.navigate("https://vip.qianlima.com/")
    gate, _ = page_gate(client)
    return gate == "ok"

def switch_to_detail_tab(client: CDPClient) -> bool:
    for selector in (".tab-bar_item", "[role=tab]", ".el-tabs__item", "text=招标详情"):
        try:
            locator = client.page.locator(selector)
            for index in range(min(locator.count(), 20)):
                item = locator.nth(index)
                if "招标详情" in normalize_text(item.inner_text()):
                    item.click(timeout=5000)
                    client.page.wait_for_timeout(1200)
                    return True
        except Exception:
            continue
    return False


def extract_body(client: CDPClient, title: str) -> Dict[str, Any]:
    switched = switch_to_detail_tab(client)
    candidates = client.page.evaluate(
        """() => {
          const selectors = ['.tab-body_item','.content-box','.nfw-cms-content',
            '.article-content','.detail-content','.bid-content','#content','article','main'];
          const out = [];
          for (const selector of selectors) {
            for (const el of document.querySelectorAll(selector)) {
              const style = getComputedStyle(el);
              if (style.display === 'none' || style.visibility === 'hidden') continue;
              const text = (el.innerText || '').trim();
              if (text.length >= 80) out.push({selector, text});
            }
          }
          if (!out.length) out.push({selector:'body', text:document.body.innerText || ''});
          return out;
        }"""
    )
    best = max(candidates or [], key=lambda item: len(item.get("text", "")),
               default={"text": "", "selector": ""})
    text = normalize_text(best.get("text", ""))
    for marker in BODY_STOP_MARKERS:
        position = text.find(marker)
        if position > MIN_BODY_CHARS:
            text = text[:position].strip()
    identity_ok = body_matches_project(title, text)
    meaningful = has_meaningful_text(text, MIN_BODY_CHARS)
    reasons = []
    if not meaningful:
        reasons.append(f"正文有效字符不足{MIN_BODY_CHARS}")
    if not identity_ok:
        reasons.append("正文与项目标题身份不匹配")
    return {
        "body_read_ok": meaningful and identity_ok,
        "body_identity_ok": identity_ok,
        "body_text": text,
        "body_text_length": len(compact_text(text)),
        "body_source": f"{'招标详情tab' if switched else '当前页'}:{best.get('selector', '')}",
        "body_invalid_reason": "；".join(reasons),
    }


TIMELINE_NODE_TYPES = (
    "\u62db\u6807\u516c\u544a",
    "\u91c7\u8d2d\u516c\u544a",
    "\u8be2\u4ef7\u516c\u544a",
    "\u7ade\u4e89\u6027\u78cb\u5546\u516c\u544a",
    "\u62db\u6807\u9884\u544a",
    "\u53d8\u66f4\u516c\u544a",
    "\u7b54\u7591\u516c\u544a",
    "\u5019\u9009\u4eba\u516c\u793a",
    "\u4e2d\u6807\u901a\u77e5",
    "\u4e2d\u6807\u7ed3\u679c",
    "\u6210\u4ea4\u516c\u544a",
    "\u5408\u540c\u516c\u544a",
)


def normalize_timeline_date(value: str) -> str:
    digits = re.findall(r"\d+", str(value or ""))
    if len(digits) >= 3:
        return f"{int(digits[0]):04d}-{int(digits[1]):02d}-{int(digits[2]):02d}"
    return normalize_text(value)


def extract_progress_timeline(client: CDPClient, project_title: str) -> List[Dict[str, str]]:
    """Read the actual Qianlima progress section; never invent template history."""
    page_text = normalize_text(client.visible_text())
    positions = [
        position
        for marker in ("\u62db\u6807\u8fdb\u5ea6\u8ddf\u8e2a", "\u8fdb\u5ea6\u8ddf\u8e2a")
        if (position := page_text.find(marker)) >= 0
    ]
    if not positions:
        return []

    segment = page_text[min(positions):]
    for marker in (
        "\u4f01\u4e1a\u5546\u60c5\u5206\u6790",
        "\u76f8\u5173\u62db\u6807",
        "\u70ed\u95e8\u63a8\u8350",
    ):
        position = segment.find(marker)
        if position > 0:
            segment = segment[:position]

    node_pattern = "|".join(re.escape(value) for value in TIMELINE_NODE_TYPES)
    date_pattern = r"20\d{2}(?:[-./]\d{1,2}[-./]\d{1,2}|\u5e74\d{1,2}\u6708\d{1,2}\u65e5)"
    pattern = re.compile(
        rf"(?P<date>{date_pattern})\s*(?:\u5f53\u524d\u4fe1\u606f)?\s*"
        rf"(?P<node>{node_pattern})\s*(?P<title>.{{4,180}}?)"
        rf"(?=(?:{date_pattern})\s*(?:\u5f53\u524d\u4fe1\u606f)?\s*(?:{node_pattern})|$)"
    )

    output = []
    seen = set()
    for match in pattern.finditer(segment):
        title = normalize_text(match.group("title"))
        title = re.split(
            r"(?:\u4f01\u4e1a\u4fe1\u606f|\+\s*\u8ddf\u8e2a|\u5546\u673a\u63a8\u8350)",
            title,
            maxsplit=1,
        )[0].strip()
        if not title:
            title = project_title
        row = {
            "date": normalize_timeline_date(match.group("date")),
            "node": normalize_text(match.group("node")),
            "title": title[:180],
        }
        key = (row["date"], row["node"], row["title"])
        if key not in seen:
            seen.add(key)
            output.append(row)
    return output


def valid_attachment_candidate(name: str, href: str) -> bool:
    clean_name = normalize_text(name)
    clean_href = normalize_text(href)
    compact_name = compact_text(clean_name)
    if not clean_href or clean_href.lower().startswith("javascript:"):
        return False
    if any(token.lower() in compact_name.lower() for token in NON_EVIDENCE_ACTION_NAMES):
        return False
    if "downloads/agent.jsp" in clean_href.lower():
        match = re.search(r"[?&]req=([^&#\s]+)", clean_href, re.I)
        if not match or not match.group(1).strip():
            return False
    suffix = Path(urlparse(clean_name).path).suffix.lower()
    if suffix in SUPPORTED_DOCUMENT_SUFFIXES:
        return True
    return any(hint in clean_name for hint in ATTACHMENT_NAME_HINTS) and compact_name not in {"附件", "下载"}


def detect_attachments(client: CDPClient) -> List[Dict[str, Any]]:
    """发现附件。

    2026-07-14 U2：优先 `.fileItem[data-filepath-id]`（VIP 真下载链路），
    agent.jsp 多附件常共用同一 req，真正区分靠 fileValidate / filepathId。
    """
    # 1) VIP fileItem（最可靠）
    file_items = client.page.evaluate(
        """() => {
          const rows = [];
          document.querySelectorAll('.fileItem, [class*="fileItem"]').forEach(item => {
            const filepathId = item.getAttribute('data-filepath-id')
              || (item.dataset && (item.dataset.filepathId || item.dataset.filepathid))
              || '';
            const fileName = item.getAttribute('data-file-name')
              || (item.dataset && (item.dataset.fileName || item.dataset.filename))
              || '';
            const nameEl = item.querySelector('.fileName, .name, a, span');
            const textName = (fileName || (nameEl ? nameEl.innerText : '') || item.innerText || '')
              .split('\\n')[0].trim();
            const down = [...item.querySelectorAll('a,button,span,div')]
              .find(x => (x.innerText || '').trim() === '下载');
            rows.push({
              name: textName,
              filepath_id: String(filepathId || ''),
              href: down && down.href ? down.href : '',
              has_download_btn: !!down,
              source: 'fileItem'
            });
          });
          return rows;
        }"""
    ) or []

    # 2) 兼容旧 DOM：agent.jsp / 附件列表链接（含 fileValidate 序号）
    raw = client.page.evaluate(
        """() => {
          const selectors = [
            '.nfw-cms-attachment a',
            '[class*="attachment-list"] a', '[class*="annex"] a',
            'a[href*="downloads/agent.jsp?req="]'
          ];
          const rows = [];
          const seen = new Set();
          for (const selector of selectors) {
            for (const el of document.querySelectorAll(selector)) {
              const link = el.matches('a') ? el : el.querySelector('a');
              const target = link || el;
              const name = (target.innerText || el.innerText || target.getAttribute('download') || '附件')
                .trim().split('\\n')[0];
              const href = target.href || target.getAttribute('href') || target.dataset.url || '';
              const onclick = target.getAttribute('onclick') || '';
              const key = name + '|' + href + '|' + onclick;
              if (seen.has(key)) continue;
              seen.add(key);
              let file_index = '';
              let file_group = '';
              const m = onclick.match(/fileValidate\\s*\\(\\s*[\"']([^\"']+)[\"']\\s*,\\s*[\"']([^\"']+)[\"']/);
              if (m) { file_group = m[1]; file_index = m[2]; }
              rows.push({
                name, href, onclick, file_group, file_index, source: 'anchor'
              });
            }
          }
          return rows;
        }"""
    ) or []

    output: List[Dict[str, Any]] = []
    seen_fid: set = set()
    seen_name: set = set()

    for item in file_items:
        name = clean_attachment_display_name(item.get("name") or "附件")
        fid = normalize_text(item.get("filepath_id"))
        if not fid:
            continue
        if fid in seen_fid:
            continue
        # 名称过滤：排除明显非附件
        compact_name = compact_text(name)
        if any(token.lower() in compact_name.lower() for token in NON_EVIDENCE_ACTION_NAMES):
            continue
        if compact_name in {"附件", "下载", "预览"}:
            continue
        seen_fid.add(fid)
        seen_name.add(compact_text(name).lower())
        output.append({
            "name": name,
            "href": normalize_text(item.get("href")),
            "filepath_id": fid,
            "has_download_btn": bool(item.get("has_download_btn")),
            "source": "fileItem",
        })

    for item in raw:
        href = normalize_text(item.get("href"))
        onclick = normalize_text(item.get("onclick"))
        if not href:
            match = re.search(
                r"https?://[^\x22\x27\s)]+|/downloads/agent\.jsp\?[^\x22\x27\s)]+", onclick
            )
            href = match.group(0) if match else ""
        name = clean_attachment_display_name(item.get("name") or "附件")
        if not valid_attachment_candidate(name, href):
            # 无 href 但有 fileValidate 的也可能有效：放宽
            if not (item.get("file_index") and item.get("file_group")):
                continue
        name_key = compact_text(name).lower()
        if name_key in seen_name:
            # 已有同名 fileItem，跳过重复 agent 链接
            continue
        # 尝试用名称回填 filepath_id（fileItem 名称子串匹配）
        filepath_id = ""
        for fi in file_items:
            fi_name = clean_attachment_display_name(fi.get("name") or "")
            if not fi.get("filepath_id"):
                continue
            if fi_name and (fi_name in name or name in fi_name or compact_text(fi_name) in compact_text(name)):
                filepath_id = str(fi.get("filepath_id"))
                break
        if filepath_id and filepath_id in seen_fid:
            continue
        seen_name.add(name_key)
        if filepath_id:
            seen_fid.add(filepath_id)
        output.append({
            "name": name,
            "href": href,
            "filepath_id": filepath_id,
            "file_group": normalize_text(item.get("file_group")),
            "file_index": normalize_text(item.get("file_index")),
            "onclick": onclick,
            "source": item.get("source") or "anchor",
        })
    return output

def request_session(client: CDPClient, referer: str) -> requests.Session:
    session = requests.Session()
    session.headers.update({"User-Agent": client.page.evaluate("() => navigator.userAgent"), "Referer": referer})
    session.cookies.update(client.cookies())
    return session


def is_static_resource_url(value: str) -> bool:
    parsed = urlparse(value)
    path = parsed.path.lower()
    host = parsed.netloc.lower()
    return (
        host.startswith("gw-static.")
        or "/css/" in path
        or "/js/" in path
        or Path(path).suffix.lower() in STATIC_RESOURCE_SUFFIXES
    )


def download_candidate_score(value: str) -> int:
    if not value.startswith(("http://", "https://")) or is_static_resource_url(value):
        return -1
    path = urlparse(value).path.lower()
    suffix = Path(path).suffix.lower()
    if suffix in SUPPORTED_DOCUMENT_SUFFIXES:
        return 100
    if "download" in path or "attachment" in path:
        return 30
    return 10


def extract_real_download_url(payload: str, base_url: str) -> str:
    decoded = html.unescape(payload).replace(chr(92) + "/", "/")
    candidates = []
    patterns = (
        r"""href\s*=\s*["']([^"']+)["']""",
        r"""https?://[^\s"'<>]+""",
    )
    for pattern in patterns:
        for candidate in re.findall(pattern, decoded, flags=re.I):
            candidate = candidate.strip().rstrip("),;]")
            if not candidate or candidate.lower().startswith("javascript:"):
                continue
            resolved = urljoin(base_url, candidate)
            if "agent.jsp" in resolved.lower():
                continue
            score = download_candidate_score(resolved)
            if score >= 0:
                candidates.append((score, resolved))
    if not candidates:
        return ""
    candidates.sort(key=lambda item: item[0], reverse=True)
    return candidates[0][1]

def validate_download_bytes(content: bytes, file_type: str, content_type: str) -> Tuple[bool, str]:
    if not content:
        return False, "下载内容为空"
    if len(content) > MAX_FILE_BYTES:
        return False, f"附件超过{MAX_FILE_BYTES // 1024 // 1024}MB限制"
    lowered_type = (content_type or "").lower()
    if lowered_type.startswith(("text/css", "text/javascript", "application/javascript")):
        return False, f"响应是静态资源而不是附件: {content_type}"
    if b"<html" in content[:500].lower() or (
        "html" in lowered_type and "xml" not in lowered_type
    ):
        return False, "响应仍是HTML跳转页"
    resolved = file_type
    if resolved == "unknown":
        resolved = sniff_file_type(content)
        if resolved == "unknown":
            return False, "无法确认附件文件类型"
    prefix = content[:8]
    if resolved == "pdf" and not prefix.startswith(b"%PDF"):
        return False, "响应不是有效PDF"
    if resolved in {"docx", "xlsx", "zip"} and not prefix.startswith(b"PK"):
        return False, f"响应不是有效{resolved.upper()}压缩容器"
    if resolved in {"doc", "xls"} and not prefix.startswith(bytes.fromhex("D0CF11E0A1B11AE1")):
        return False, f"响应不是有效{resolved.upper()}复合文档"
    return True, ""

def browser_capture_spa_download(client: CDPClient, url: str) -> Dict[str, Any]:
    """Resolve SPA download routes with a bounded wait and explicit captcha gate."""
    page = client.context.new_page()
    page.set_default_timeout(15_000)
    responses = []
    downloads = []
    extra_pages = []

    def remember_response(response: Any) -> None:
        try:
            headers = response.all_headers()
            content_type = headers.get("content-type", "")
            disposition = headers.get("content-disposition", "")
            candidate = response.url
            score = download_candidate_score(candidate)
            if disposition or score >= 30 or any(
                token in content_type.lower()
                for token in ("pdf", "word", "excel", "spreadsheet", "octet-stream")
            ):
                responses.append((response, content_type, disposition))
        except Exception:
            pass

    def remember_popup(popup: Any) -> None:
        # 2026-07-13 新标签页：绑定相同监听器，并尝试点击其中的中间页按钮
        try:
            extra_pages.append(popup)
            popup.on("response", remember_response)
            popup.on("download", lambda d: downloads.append(d))
            popup.on("popup", remember_popup)
            try:
                popup.wait_for_load_state("domcontentloaded", timeout=8_000)
            except Exception:
                pass
            try:
                ch = popup.locator("text=点击此处")
                if ch.count() > 0 and ch.first.is_visible():
                    ch.first.click(timeout=3000)
            except Exception:
                pass
        except Exception:
            pass

    page.on("response", remember_response)
    page.on("download", lambda download: downloads.append(download))
    page.on("popup", remember_popup)
    try:
        try:
            page.goto(url, wait_until="domcontentloaded", timeout=20_000)
        except Exception as exc:
            return {"ok": False, "error": f"SPA页面打开失败: {exc}"}

        # 2026-07-10 修复：千里马 agent.jsp 中间页面需要点击"点击此处"按钮
        # 2026-07-13："点击此处"的 href 直接就是外部站真实下载地址，
        # 先提取并直接导航检测状态码，区分死链（404）与真超时。
        real_link = ""
        try:
            hrefs = page.eval_on_selector_all(
                "a",
                "els => els.map(e => ({text: e.innerText, href: e.href}))",
            )
            for lk in hrefs:
                if "点击此处" in (lk.get("text") or ""):
                    real_link = lk.get("href", "")
                    break
        except Exception:
            pass

        # 直接导航到真实链接：成功则触发下载/响应；失败则拿到准确状态码
        if real_link and not real_link.lower().startswith("javascript:"):
            try:
                nav_resp = page.goto(real_link, timeout=15_000)
                if nav_resp is not None and nav_resp.status >= 400:
                    return {"ok": False,
                            "error": f"外部附件源文件不可用（HTTP {nav_resp.status}，可能已从源站删除）：{real_link[:120]}"}
            except Exception:
                # 导航异常通常是下载中断导航，交给下方轮询捕获
                pass

        # 先检查是否有下载按钮，如果有就点击
        try:
            click_here = page.locator("text=点击此处")
            if click_here.count() > 0 and click_here.first.is_visible():
                click_here.first.click(timeout=3000)
                page.wait_for_timeout(1000)
        except Exception:
            pass

        # 2026-07-14 P0：SPA 轮询有 wall 预算 + 日志，不再无声 30s
        spa_start = time.time()
        last_signal_at = spa_start
        last_resp_n = 0
        last_dl_n = 0
        poll_i = 0
        slog(f"  [SPA] poll start budget={SPA_POLL_MAX_SEC}s url={url[:100]}")
        while (time.time() - spa_start) < SPA_POLL_MAX_SEC:
            poll_i += 1
            page.wait_for_timeout(1_000)
            try:
                visible = page.locator("body").inner_text(timeout=2_000)
            except Exception:
                visible = ""
            if any(token in visible for token in ("验证码", "看不清", "安全验证", "访问验证")):
                slog("  [SPA] captcha detected")
                return {"ok": False, "error": "外部附件站点要求验证码，不能无人值守读取，需人工复核"}
            for extra in list(extra_pages):
                try:
                    ch = extra.locator("text=点击此处")
                    if ch.count() > 0 and ch.first.is_visible():
                        ch.first.click(timeout=3000)
                        last_signal_at = time.time()
                except Exception:
                    pass
            if len(responses) > last_resp_n or len(downloads) > last_dl_n:
                last_resp_n = len(responses)
                last_dl_n = len(downloads)
                last_signal_at = time.time()
                slog(f"  [SPA] signal downloads={last_dl_n} responses={last_resp_n} t={int(time.time()-spa_start)}s")
            if downloads:
                download = downloads[-1]
                try:
                    import tempfile
                    temp_download = Path(tempfile.gettempdir()) / f"cdp_download_{time.time()}.tmp"
                    download.save_as(str(temp_download))
                    if temp_download.exists():
                        content = temp_download.read_bytes()
                        temp_download.unlink()
                        # U2：空/HTML 不当作成功
                        if content and b"<html" not in content[:400].lower():
                            slog(f"  [SPA] browser_download_event bytes={len(content)}")
                            return {
                                "ok": True,
                                "content": content,
                                "content_type": "application/octet-stream",
                                "real_url": download.url,
                                "suggested_filename": download.suggested_filename,
                                "method": "browser_download_event",
                            }
                except Exception as e:
                    slog(f"  [SPA] download wait failed: {e}")
            # U2：有 response 时先尝试验证 body，无效则继续轮询（勿 1s 假失败）
            if responses:
                for response, content_type, _ in reversed(responses):
                    try:
                        content_length = int(
                            (response.all_headers().get("content-length") or "0").strip() or 0
                        )
                        if content_length > MAX_FILE_BYTES:
                            continue
                        content = response.body()
                        if not content or len(content) > MAX_FILE_BYTES:
                            continue
                        if b"<html" in content[:500].lower() or (
                            "html" in (content_type or "").lower() and "xml" not in (content_type or "").lower()
                        ):
                            continue
                        if sniff_file_type(content) == "unknown" and len(content) < 64:
                            continue
                        slog(f"  [SPA] browser_network_response bytes={len(content)}")
                        return {
                            "ok": True,
                            "content": content,
                            "content_type": content_type,
                            "real_url": response.url,
                            "suggested_filename": "",
                            "method": "browser_network_response",
                        }
                    except Exception:
                        continue
            if (time.time() - last_signal_at) >= max(20, SPA_POLL_MAX_SEC // 2) and poll_i >= 10:
                slog(f"  [SPA] stall no signal for {int(time.time()-last_signal_at)}s")
                break
            if poll_i % 10 == 0:
                slog(f"  [SPA] heartbeat poll={poll_i} elapsed={int(time.time()-spa_start)}s")

        slog(f"  [SPA] fail after {int(time.time()-spa_start)}s")
        if real_link:
            return {"ok": False,
                    "error": f"外部附件未捕获到下载响应（已尝试真实链接 {real_link[:100]}，可能需人工复核）"}
        return {"ok": False, "error": "SPA页面未捕获到真实附件下载响应（未找到中间页链接）"}
    finally:
        try:
            page.close()
        except Exception:
            pass


def stream_download_to_file(
    url: str,
    target_path: Path,
    session: Optional[requests.Session] = None,
    timeout: Optional[Tuple[int, int]] = None,
) -> Dict[str, Any]:
    """流式下载；2026-07-14 P0：读超时收紧 + 字节增长看门狗。"""
    if timeout is None:
        timeout = (DOWNLOAD_CONNECT_TIMEOUT_SEC, DOWNLOAD_READ_TIMEOUT_SEC)
    temp_path = target_path.with_suffix(target_path.suffix + ".part")
    req_session = session or requests.Session()
    start = time.time()
    last_size = 0
    last_grow_at = start
    last_hb = start

    slog(
        f"  [DL] stream start timeout={timeout} stall={DOWNLOAD_STALL_SEC}s "
        f"hard={DOWNLOAD_HARD_CAP_SEC}s url={url[:120]}"
    )
    try:
        with req_session.get(url, stream=True, timeout=timeout) as response:
            response.raise_for_status()
            total_size = 0
            with open(temp_path, "wb") as f:
                for chunk in response.iter_content(chunk_size=1024 * 1024):
                    now = time.time()
                    if chunk:
                        f.write(chunk)
                        total_size += len(chunk)
                        if total_size > last_size:
                            last_size = total_size
                            last_grow_at = now
                    if total_size > MAX_FILE_BYTES:
                        raise RuntimeError(f"download exceeds MAX_FILE_BYTES ({MAX_FILE_BYTES})")
                    if (now - start) >= DOWNLOAD_HARD_CAP_SEC:
                        raise RuntimeError(
                            f"download hard_cap {DOWNLOAD_HARD_CAP_SEC}s size={total_size}"
                        )
                    if (now - last_grow_at) >= DOWNLOAD_STALL_SEC and total_size > 0:
                        raise RuntimeError(
                            f"download stalled {DOWNLOAD_STALL_SEC}s at size={total_size}"
                        )
                    if (now - last_grow_at) >= DOWNLOAD_STALL_SEC and total_size == 0 and (now - start) >= DOWNLOAD_STALL_SEC:
                        raise RuntimeError(f"download no bytes for {DOWNLOAD_STALL_SEC}s")
                    if (now - last_hb) >= 30:
                        slog(
                            f"  [DL] heartbeat size={total_size} "
                            f"elapsed={int(now-start)}s stall_for={int(now-last_grow_at)}s"
                        )
                        last_hb = now

            if total_size < 100:
                slog(f"  [DL] fail too_small size={total_size}")
                return {"ok": False, "error": "下载文件过小", "size": total_size}

            if target_path.exists():
                try:
                    target_path.unlink()
                except Exception:
                    pass
            temp_path.replace(target_path)
            slog(f"  [DL] ok size={total_size} elapsed={int(time.time()-start)}s")
            return {"ok": True, "size": total_size, "error": ""}

    except Exception as exc:
        slog(f"  [DL] fail: {exc}")
        try:
            if temp_path.exists():
                temp_path.unlink()
        except Exception:
            pass
        return {"ok": False, "error": f"流式下载失败: {exc}", "size": 0}


def download_via_preview_api(client: CDPClient, filepath_id: str, attachment_name: str, page_url: str,
                             target_dir: Path, index: int) -> Dict[str, Any]:
    """2026-07-10 新增：通过千里马预览API获取PDF数据。
    
    流程：
    1. 从页面URL提取contentId
    2. 调用API获取previewUrl
    3. 用浏览器打开预览页（需要浏览器环境）
    4. 等待PDF.js加载，提取PDF数据
    5. 保存为PDF文件
    
    注意：大文件（>20MB）的PDF数据可能以base64嵌入HTML，需要等待页面完全加载。
    """
    import base64
    
    if not filepath_id:
        return {"ok": False, "error": "缺少filepathId"}
    
    # 从页面URL提取contentId
    bid_match = re.search(r"bid-(\d+)", page_url)
    if not bid_match:
        return {"ok": False, "error": "无法从URL提取contentId"}
    content_id = bid_match.group(1)
    
    # 调用API获取previewUrl
    try:
        api_result = client.page.evaluate(
            """async ([contentId, filepathId]) => {
                try {
                    const url = `/rest/detail/alltypesdetail/getPreViewByPathId?contentId=${contentId}&filepathId=${filepathId}`;
                    const response = await fetch(url, {
                        method: 'GET',
                        credentials: 'include',
                        headers: {
                            'Accept': 'application/json',
                            'X-Requested-With': 'XMLHttpRequest'
                        }
                    });
                    if (!response.ok) {
                        return { ok: false, status: response.status };
                    }
                    const data = await response.json();
                    return { ok: true, data: data };
                } catch (err) {
                    return { ok: false, error: err.message };
                }
            }""",
            [content_id, filepath_id]
        )
    except Exception as exc:
        return {"ok": False, "error": f"API调用失败: {exc}"}
    
    if not api_result.get("ok"):
        return {"ok": False, "error": f"API返回失败: {api_result.get('error', 'unknown')}"}
    
    # 解析API返回数据
    data = api_result.get("data", {})
    inner_data = data.get("data", data)
    preview_url = inner_data.get("previewUrl") or inner_data.get("url")
    api_file_name = inner_data.get("fileName", attachment_name)
    
    if not preview_url:
        return {"ok": False, "error": "API未返回previewUrl"}
    
    # 用浏览器打开预览页（需要浏览器环境，requests会返回400）
    preview_page = client.context.new_page()
    preview_page.set_default_timeout(120_000)  # 增加超时到120秒
    
    try:
        # 不等待domcontentloaded，大文件加载很慢
        try:
            preview_page.goto(preview_url, wait_until="commit", timeout=60_000)
        except Exception as exc:
            # 即使超时，页面可能已部分加载，继续尝试提取
            pass
        
        # 2026-07-14 P0：PDF.js 等待有 wall 预算 + stall，不再无声 90s
        pdf_ready = False
        preview_start = time.time()
        last_preview_signal = preview_start
        last_marker = ""
        slog(f"  [PREVIEW] wait pdf.js budget={PREVIEW_POLL_MAX_SEC}s stall={PREVIEW_STALL_SEC}s")
        while (time.time() - preview_start) < PREVIEW_POLL_MAX_SEC:
            preview_page.wait_for_timeout(1000)
            try:
                check = preview_page.evaluate("""() => {
                    if (window.PDFViewerApplication && window.PDFViewerApplication.pdfDocument) {
                        return { ready: true, numPages: window.PDFViewerApplication.pdfDocument.numPages };
                    }
                    // 检查xdocViewOption.data（base64嵌入）
                    if (window.xdocViewOption && window.xdocViewOption.data) {
                        return { ready: true, base64: true, size: window.xdocViewOption.data.length };
                    }
                    return { ready: false };
                }""")
                marker = json.dumps(check, ensure_ascii=False, sort_keys=True) if isinstance(check, dict) else str(check)
                if marker != last_marker:
                    last_marker = marker
                    last_preview_signal = time.time()
                if check.get("ready"):
                    pdf_ready = True
                    slog(
                        f"  [PREVIEW] ready elapsed={int(time.time()-preview_start)}s "
                        f"info={str(check)[:120]}"
                    )
                    break
            except Exception:
                pass
            if (time.time() - last_preview_signal) >= PREVIEW_STALL_SEC:
                slog(f"  [PREVIEW] stall no state change for {PREVIEW_STALL_SEC}s")
                break
            if int(time.time() - preview_start) % 15 == 0:
                slog(f"  [PREVIEW] heartbeat elapsed={int(time.time()-preview_start)}s")

        if not pdf_ready:
            # 最后尝试从HTML源码提取base64数据
            try:
                html_content = preview_page.content()
                data_match = re.search(r'xdocViewOption\.data\s*=\s*"([^"]+)"', html_content)
                if data_match:
                    b64_data = data_match.group(1)
                    pdf_bytes = base64.b64decode(b64_data)
                    if pdf_bytes[:4].startswith(b'%PDF'):
                        name = safe_filename(api_file_name or attachment_name or f"attachment_{index}")
                        if not Path(name).suffix:
                            name += ".pdf"
                        target_dir.mkdir(parents=True, exist_ok=True)
                        target_path = target_dir / f"{index:02d}_{name}"
                        target_path.write_bytes(pdf_bytes)
                        return {
                            "ok": True,
                            "download_path": str(target_path),
                            "size": len(pdf_bytes),
                            "real_url": preview_url,
                            "preview_url": preview_url,
                            "suggested_filename": api_file_name or attachment_name,
                            "method": "preview_api_html_base64",
                        }
            except Exception:
                pass
            slog(f"  [PREVIEW] timeout/stall after {int(time.time()-preview_start)}s")
            return {
                "ok": False,
                "error": f"PDF.js加载超时/停滞（预算{PREVIEW_POLL_MAX_SEC}s，stall{PREVIEW_STALL_SEC}s）",
            }
        
        # 尝试从PDF.js提取数据
        pdf_result = preview_page.evaluate(
            """async () => {
                try {
                    // 方法1: 从PDFViewerApplication获取数据
                    if (window.PDFViewerApplication && window.PDFViewerApplication.pdfDocument) {
                        const pdfDoc = window.PDFViewerApplication.pdfDocument;
                        if (pdfDoc.getData) {
                            const data = await pdfDoc.getData();
                            if (data && data.byteLength > 0) {
                                const bytes = new Uint8Array(data);
                                let binary = '';
                                for (let i = 0; i < bytes.length; i++) {
                                    binary += String.fromCharCode(bytes[i]);
                                }
                                return {
                                    ok: true,
                                    base64: btoa(binary),
                                    size: data.byteLength,
                                    pageCount: pdfDoc.numPages || 0,
                                    method: 'pdfjs_getData'
                                };
                            }
                        }
                    }
                    // 方法2: 从xdocViewOption获取（base64嵌入）
                    if (window.xdocViewOption && window.xdocViewOption.data) {
                        return {
                            ok: true,
                            base64: window.xdocViewOption.data,
                            size: window.xdocViewOption.data.length * 3 / 4,
                            pageCount: 0,
                            method: 'xdocViewOption'
                        };
                    }
                    return { ok: false, error: "PDFViewerApplication和xdocViewOption均不可用" };
                } catch (e) {
                    return { ok: false, error: e.message };
                }
            }"""
        )
        
        if not pdf_result.get("ok"):
            return {"ok": False, "error": f"PDF数据提取失败: {pdf_result.get('error', 'unknown')}"}
        
        # 解码base64数据
        b64_data = pdf_result.get("base64", "")
        if not b64_data:
            return {"ok": False, "error": "PDF数据为空"}
        
        try:
            pdf_bytes = base64.b64decode(b64_data)
        except Exception as exc:
            return {"ok": False, "error": f"base64解码失败: {exc}"}
        
        if not pdf_bytes or len(pdf_bytes) < 100:
            return {"ok": False, "error": "PDF数据过小"}
        
        # 验证PDF头
        if not pdf_bytes[:4].startswith(b'%PDF'):
            return {"ok": False, "error": f"数据不是有效PDF，头部: {pdf_bytes[:10]}"}
        
        # 保存文件
        name = safe_filename(api_file_name or attachment_name or f"attachment_{index}")
        if not Path(name).suffix:
            name += ".pdf"
        target_dir.mkdir(parents=True, exist_ok=True)
        target_path = target_dir / f"{index:02d}_{name}"
        target_path.write_bytes(pdf_bytes)
        
        return {
            "ok": True,
            "download_path": str(target_path),
            "size": len(pdf_bytes),
            "real_url": preview_url,
            "preview_url": preview_url,
            "suggested_filename": api_file_name or attachment_name,
            "method": pdf_result.get("method", "preview_api"),
            "page_count": pdf_result.get("pageCount", 0),
        }
    
    except Exception as exc:
        return {"ok": False, "error": f"预览页面异常: {exc}"}
    finally:
        try:
            preview_page.close()
        except Exception:
            pass


def _finalize_downloaded_bytes(
    content: bytes,
    *,
    name_hint: str,
    content_type: str,
    real_url: str,
    source_url: str,
    target_dir: Path,
    index: int,
    capture_method: str,
) -> Dict[str, Any]:
    """校验字节 → 写盘。U2：unknown 时用魔数 sniff。"""
    name = safe_filename(name_hint or f"attachment_{index}")
    file_type = infer_file_type(name, content_type)
    if file_type == "unknown":
        file_type = infer_file_type(real_url, content_type)
    if file_type == "unknown":
        file_type = sniff_file_type(content)
    valid, reason = validate_download_bytes(content, file_type, content_type)
    if not valid:
        return {
            "download_ok": False,
            "error": reason,
            "source_url": source_url,
            "real_url": real_url,
            "content_type": content_type,
            "file_type": file_type,
        }
    if file_type == "unknown":
        file_type = sniff_file_type(content)
    if not Path(name).suffix:
        name += {
            "docx": ".docx", "doc": ".doc", "pdf": ".pdf",
            "xlsx": ".xlsx", "xls": ".xls", "zip": ".zip",
            "rar": ".rar", "7z": ".7z", "image": ".bin",
        }.get(file_type, ".bin")
    digest = hashlib.sha256(content).hexdigest()
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / f"{index:02d}_{digest[:8]}_{name}"
    target.write_bytes(content)
    return {
        "download_ok": True,
        "download_path": str(target),
        "sha256": digest,
        "source_url": source_url,
        "real_url": real_url,
        "content_type": content_type,
        "file_type": file_type,
        "bytes": len(content),
        "capture_method": capture_method,
        "error": "",
    }


def download_via_filepath_stream(
    client: CDPClient,
    content_id: str,
    filepath_id: str,
    attachment_name: str,
    target_dir: Path,
    index: int,
) -> Dict[str, Any]:
    """2026-07-14 U2：VIP 附件真下载。

    链路（已在实机验证）：
      getFileStreamPreCheckPermission → getZBFileStreamByPathId(302) → OSS 文件
    使用 Playwright context.request（带浏览器 Cookie），避免 CORS。
    """
    if not content_id or not filepath_id:
        return {"ok": False, "error": "缺少 contentId/filepathId"}
    base = "https://detail.vip.qianlima.com/rest/detail/alltypesdetail"
    pre_url = (
        f"{base}/getFileStreamPreCheckPermission"
        f"?contentId={content_id}&filepathId={filepath_id}&_={int(time.time() * 1000)}"
    )
    stream_url = (
        f"{base}/getZBFileStreamByPathId"
        f"?contentId={content_id}&filepathId={filepath_id}"
    )
    slog(f"  [STREAM] precheck filepathId={filepath_id}")
    try:
        api = client.context.request
        pre = api.get(pre_url, timeout=30_000)
        if pre.status >= 400:
            return {"ok": False, "error": f"预检HTTP {pre.status}"}
        try:
            pre_json = pre.json()
        except Exception:
            pre_json = {}
        # code 200 / data true 为常见成功形态；不强制，避免站点改字段即全挂
        code = pre_json.get("code")
        if code not in (None, 200, "200") and pre_json.get("data") is False:
            return {
                "ok": False,
                "error": f"预检拒绝: code={code} msg={pre_json.get('msg')}",
            }
        slog(f"  [STREAM] getZBFileStreamByPathId filepathId={filepath_id}")
        resp = api.get(stream_url, timeout=120_000, max_redirects=15)
        if resp.status >= 400:
            return {"ok": False, "error": f"流式下载HTTP {resp.status}"}
        content = resp.body()
        if not content:
            return {"ok": False, "error": "流式下载内容为空"}
        content_type = resp.headers.get("content-type") or ""
        final_url = resp.url or stream_url
        # HTML 错误页
        if b"<html" in content[:400].lower() and not content.startswith(b"%PDF"):
            return {"ok": False, "error": f"流式下载返回HTML而非文件 ct={content_type}"}
        slog(f"  [STREAM] ok bytes={len(content)} ct={content_type[:40]} url={final_url[:80]}")
        finalized = _finalize_downloaded_bytes(
            content,
            name_hint=attachment_name,
            content_type=content_type,
            real_url=final_url,
            source_url=stream_url,
            target_dir=target_dir,
            index=index,
            capture_method="zb_file_stream",
        )
        if finalized.get("download_ok"):
            return {"ok": True, **finalized}
        return {"ok": False, "error": finalized.get("error") or "流式下载校验失败", **finalized}
    except Exception as exc:
        return {"ok": False, "error": f"流式下载异常: {exc}"}


def download_via_fileitem_click(
    client: CDPClient,
    filepath_id: str,
    attachment_name: str,
    target_dir: Path,
    index: int,
) -> Dict[str, Any]:
    """回退：在详情页点击 fileItem「下载」按钮，捕获 Playwright download 事件。"""
    if not filepath_id:
        return {"ok": False, "error": "缺少 filepathId"}
    slog(f"  [CLICK] fileItem download filepathId={filepath_id}")
    try:
        with client.page.expect_download(timeout=90_000) as download_info:
            clicked = client.page.evaluate(
                """(fid) => {
                  const items = [...document.querySelectorAll('.fileItem, [class*="fileItem"]')];
                  const it = items.find(x =>
                    (x.getAttribute('data-filepath-id') || x.dataset.filepathId || '') === String(fid)
                  );
                  if (!it) return {ok: false, reason: 'fileItem_not_found'};
                  const btn = [...it.querySelectorAll('a,button,span,div')]
                    .find(x => (x.innerText || '').trim() === '下载');
                  if (btn) { btn.click(); return {ok: true, via: 'download_btn'}; }
                  it.click();
                  return {ok: true, via: 'item_click'};
                }""",
                str(filepath_id),
            )
            if not (clicked or {}).get("ok"):
                raise RuntimeError(f"未找到可点击的 fileItem: {clicked}")
        download = download_info.value
        import tempfile
        temp_download = Path(tempfile.gettempdir()) / f"cdp_fileitem_{int(time.time() * 1000)}.tmp"
        download.save_as(str(temp_download))
        content = temp_download.read_bytes()
        try:
            temp_download.unlink()
        except Exception:
            pass
        suggested = download.suggested_filename or attachment_name
        finalized = _finalize_downloaded_bytes(
            content,
            name_hint=suggested or attachment_name,
            content_type="application/octet-stream",
            real_url=download.url or "",
            source_url=f"fileItem:{filepath_id}",
            target_dir=target_dir,
            index=index,
            capture_method="fileitem_click_download",
        )
        if finalized.get("download_ok"):
            return {"ok": True, **finalized}
        return {"ok": False, "error": finalized.get("error") or "点击下载校验失败", **finalized}
    except Exception as exc:
        return {"ok": False, "error": f"fileItem点击下载失败: {exc}"}


def download_attachment(client: CDPClient, attachment: Dict[str, Any], target_dir: Path,
                        index: int, page_url: str,
                        budget_deadline: Optional[float] = None) -> Dict[str, Any]:
    href_raw = normalize_text(attachment.get("href"))
    href = urljoin(page_url, href_raw) if href_raw else ""
    att_name = clean_attachment_display_name(attachment.get("name", ""))[:80]
    filepath_id = normalize_text(attachment.get("filepath_id"))
    att_start = time.time()
    if budget_deadline is not None and time.time() >= budget_deadline:
        slog(f"  [ATTACH] skip budget exhausted before start name={att_name}")
        return {"download_ok": False, "error": "attachment budget exhausted before start"}

    slog(
        f"  [ATTACH] begin #{index} name={att_name} "
        f"filepath_id={filepath_id or '-'} href={(href or '')[:90]}"
    )

    bid_match = re.search(r"bid-(\d+)", page_url or "")
    content_id = bid_match.group(1) if bid_match else ""
    errors: List[str] = []
    # 显式初始化：下载失败 fallback 仍可回填 stream/click 拿到的 API URL（勿用 "x" in dir()）
    stream_result: Dict[str, Any] = {}
    click_result: Dict[str, Any] = {}

    # ---- U2 主路径：filepathId + VIP 流式下载 ----
    if filepath_id and content_id:
        if budget_deadline is not None and time.time() >= budget_deadline:
            return {"download_ok": False, "error": "attachment budget exhausted before stream"}
        slog(f"  [ATTACH] stage=zb_file_stream #{index}")
        t0 = time.time()
        stream_result = download_via_filepath_stream(
            client, content_id, filepath_id, att_name, target_dir, index
        )
        slog(
            f"  [ATTACH] stage=zb_file_stream done ok={stream_result.get('ok')} "
            f"dt={time.time() - t0:.1f}s"
        )
        if stream_result.get("ok") and stream_result.get("download_ok"):
            return {k: v for k, v in stream_result.items() if k != "ok"}
        if stream_result.get("error"):
            errors.append(f"stream:{stream_result.get('error')}")

        # 回退：点击 fileItem 下载
        if budget_deadline is None or time.time() < budget_deadline:
            slog(f"  [ATTACH] stage=fileitem_click #{index}")
            t1 = time.time()
            click_result = download_via_fileitem_click(
                client, filepath_id, att_name, target_dir, index
            )
            slog(
                f"  [ATTACH] stage=fileitem_click done ok={click_result.get('ok')} "
                f"dt={time.time() - t1:.1f}s"
            )
            if click_result.get("ok") and click_result.get("download_ok"):
                return {k: v for k, v in click_result.items() if k != "ok"}
            if click_result.get("error"):
                errors.append(f"click:{click_result.get('error')}")

    # ---- 预览 API（PDF 优先，且有 filepathId）----
    preview_result: Dict[str, Any] = {}
    attachment_name_hint = att_name.lower()
    is_pdf_hint = attachment_name_hint.endswith(".pdf") or (
        filepath_id
        and not any(
            attachment_name_hint.endswith(ext)
            for ext in (".doc", ".docx", ".xls", ".xlsx", ".zip", ".rar", ".7z", ".ppt", ".pptx")
        )
    )
    if filepath_id and is_pdf_hint:
        if budget_deadline is not None and time.time() >= budget_deadline:
            return {"download_ok": False, "error": "attachment budget exhausted before preview"}
        slog(f"  [ATTACH] stage=preview_api #{index}")
        t0 = time.time()
        preview_result = download_via_preview_api(
            client, filepath_id, att_name, page_url, target_dir, index
        )
        slog(f"  [ATTACH] stage=preview_api done ok={preview_result.get('ok')} dt={time.time()-t0:.1f}s")
        if preview_result.get("ok"):
            download_path = Path(preview_result["download_path"])
            content = download_path.read_bytes()
            finalized = _finalize_downloaded_bytes(
                content,
                name_hint=download_path.name or att_name,
                content_type="application/pdf",
                real_url=preview_result.get("real_url", href or page_url),
                source_url=href or page_url,
                target_dir=target_dir,
                index=index,
                capture_method=preview_result.get("method", "preview_api_stream"),
            )
            if finalized.get("download_ok"):
                if download_path.exists() and Path(finalized["download_path"]).resolve() != download_path.resolve():
                    try:
                        download_path.unlink()
                    except Exception:
                        pass
                return finalized
            try:
                download_path.unlink()
            except Exception:
                pass
            errors.append(f"preview_validate:{finalized.get('error')}")
        elif preview_result.get("error"):
            errors.append(f"preview:{preview_result.get('error')}")

    # ---- agent.jsp / 外链 SPA 回退 ----
    if href and not href.lower().startswith("javascript:"):
        if budget_deadline is not None and time.time() >= budget_deadline:
            slog(f"  [ATTACH] budget exhausted before SPA/requests #{index}")
            return {"download_ok": False, "error": "attachment budget exhausted before SPA"}
        # agent.jsp 或已确认需要浏览器的地址：SPA 捕获
        if "agent.jsp" in href.lower():
            slog(f"  [ATTACH] stage=spa_fallback #{index}")
            t_spa = time.time()
            captured = browser_capture_spa_download(client, href)
            slog(f"  [ATTACH] stage=spa_fallback done ok={captured.get('ok')} dt={time.time()-t_spa:.1f}s")
            if captured.get("ok"):
                finalized = _finalize_downloaded_bytes(
                    captured["content"],
                    name_hint=captured.get("suggested_filename") or att_name,
                    content_type=captured.get("content_type", ""),
                    real_url=captured.get("real_url", href),
                    source_url=href,
                    target_dir=target_dir,
                    index=index,
                    capture_method=captured.get("method", "browser"),
                )
                if finalized.get("download_ok"):
                    return finalized
                errors.append(f"spa_validate:{finalized.get('error')}")
            else:
                errors.append(f"spa:{captured.get('error', '浏览器下载失败')}")
        else:
            # 普通直链：requests（带 cookie）+ HTML 跳转时再 SPA
            slog(f"  [ATTACH] stage=requests #{index}")
            session = request_session(client, page_url)
            try:
                response = session.get(href, timeout=45, allow_redirects=True)
                response.raise_for_status()
                content = response.content
                content_type = response.headers.get("Content-Type", "")
                real_url = response.url
                if "html" in (content_type or "").lower() or b"<html" in content[:500].lower():
                    candidate = extract_real_download_url(response.text, response.url)
                    if candidate:
                        response = session.get(candidate, timeout=60, allow_redirects=True)
                        response.raise_for_status()
                        content = response.content
                        content_type = response.headers.get("Content-Type", "")
                        real_url = response.url
                    else:
                        captured = browser_capture_spa_download(client, real_url or href)
                        if captured.get("ok"):
                            content = captured["content"]
                            content_type = captured.get("content_type", "")
                            real_url = captured.get("real_url", real_url)
                        else:
                            errors.append(f"requests:跳转页无真实地址; spa:{captured.get('error')}")
                            content = b""
                if content:
                    finalized = _finalize_downloaded_bytes(
                        content,
                        name_hint=att_name,
                        content_type=content_type,
                        real_url=real_url,
                        source_url=href,
                        target_dir=target_dir,
                        index=index,
                        capture_method="requests",
                    )
                    if finalized.get("download_ok"):
                        return finalized
                    # HTML 误判时再 SPA
                    if finalized.get("error") == "响应仍是HTML跳转页":
                        captured = browser_capture_spa_download(client, real_url or href)
                        if captured.get("ok"):
                            finalized = _finalize_downloaded_bytes(
                                captured["content"],
                                name_hint=captured.get("suggested_filename") or att_name,
                                content_type=captured.get("content_type", ""),
                                real_url=captured.get("real_url", real_url),
                                source_url=href,
                                target_dir=target_dir,
                                index=index,
                                capture_method=captured.get("method", "browser"),
                            )
                            if finalized.get("download_ok"):
                                return finalized
                            errors.append(f"spa_validate:{finalized.get('error')}")
                        else:
                            errors.append(f"spa:{captured.get('error')}")
                    else:
                        errors.append(f"requests_validate:{finalized.get('error')}")
            except Exception as exc:
                errors.append(f"requests:{exc}")

    if not href and not filepath_id:
        return {"download_ok": False, "error": "附件缺少有效下载链接或filepathId"}

    error_msg = "；".join(errors) if errors else "附件下载失败"
    # zip 校验失败等：仍保留 stream/click 过程中的 API URL，避免 Excel 空链接挡 Phase34
    src = (
        str(stream_result.get("source_url") or "").strip()
        or str(click_result.get("source_url") or "").strip()
        or href
        or ""
    )
    real = (
        str(stream_result.get("real_url") or "").strip()
        or str(click_result.get("real_url") or "").strip()
        or src
        or href
        or ""
    )
    return {
        "download_ok": False,
        "error": error_msg,
        "source_url": src,
        "real_url": real,
    }


def parse_docx(path: Path) -> str:
    if Document is None:
        raise RuntimeError("缺少python-docx")
    document = Document(str(path))
    chunks = [p.text for p in document.paragraphs if normalize_text(p.text)]
    for table in document.tables:
        for row in table.rows:
            chunks.append(" | ".join(normalize_text(cell.text) for cell in row.cells))
    return "\n".join(chunks)


def parse_doc_ole_strings(path: Path) -> str:
    """粗提 OLE .doc 内可见 Unicode/GBK 串（无 Word 时的降级）。"""
    try:
        import olefile  # type: ignore
    except ImportError as exc:
        raise RuntimeError("缺少olefile，无法降级解析.doc") from exc
    if not olefile.isOleFile(str(path)):
        raise RuntimeError("不是OLE复合文档(.doc)")
    chunks: List[str] = []
    with olefile.OleFileIO(str(path)) as ole:
        for entry in ole.listdir():
            try:
                data = ole.openstream(entry).read()
            except Exception:
                continue
            # UTF-16LE runs
            try:
                u = data.decode("utf-16-le", errors="ignore")
                parts = re.findall(r"[\u4e00-\u9fffA-Za-z0-9，。；：、（）()\-_/ \t]{8,}", u)
                chunks.extend(parts[:40])
            except Exception:
                pass
            try:
                g = data.decode("gbk", errors="ignore")
                parts = re.findall(r"[\u4e00-\u9fffA-Za-z0-9，。；：、（）()\-_/ \t]{8,}", g)
                chunks.extend(parts[:40])
            except Exception:
                pass
    text = "\n".join(normalize_text(c) for c in chunks if normalize_text(c))
    if len(compact_text(text)) < 40:
        raise RuntimeError("OLE粗提文本不足")
    return text


def parse_doc(path: Path) -> str:
    """解析旧版 .doc：误标 docx → python-docx；否则 Word COM；再 OLE 粗提。"""
    head = path.read_bytes()[:8]
    if head[:2] == b"PK":
        return parse_docx(path)
    # Windows + 本机 Word：最可靠
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore

        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            try:
                word.DisplayAlerts = 0
            except Exception:
                pass
            # Word 需要绝对路径
            abs_path = str(path.resolve())
            doc = word.Documents.Open(abs_path, ReadOnly=True, AddToRecentFiles=False)
            text = str(doc.Content.Text or "")
            return text
        finally:
            try:
                if doc is not None:
                    doc.Close(False)
            except Exception:
                pass
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
    except Exception as com_exc:
        # 降级 OLE
        try:
            return parse_doc_ole_strings(path)
        except Exception as ole_exc:
            raise RuntimeError(f"doc解析失败 COM={com_exc}; OLE={ole_exc}") from ole_exc


def parse_pdf(path: Path) -> str:
    if pdfplumber is None:
        raise RuntimeError("缺少pdfplumber")
    with pdfplumber.open(str(path)) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def parse_xlsx(path: Path) -> str:
    if load_workbook is None:
        raise RuntimeError("缺少openpyxl")
    workbook = load_workbook(str(path), read_only=True, data_only=True)
    chunks = []
    try:
        for sheet in workbook.worksheets:
            chunks.append(f"[{sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                values = [normalize_text(value) for value in row if value not in (None, "")]
                if values:
                    chunks.append(" | ".join(values))
    finally:
        workbook.close()
    return "\n".join(chunks)


def parse_xls(path: Path) -> str:
    """旧版 xls：优先 xlrd；失败则提示不支持。"""
    try:
        import xlrd  # type: ignore
    except ImportError as exc:
        raise RuntimeError("缺少xlrd，无法解析.xls") from exc
    book = xlrd.open_workbook(str(path))
    chunks = []
    for sheet in book.sheets():
        chunks.append(f"[{sheet.name}]")
        for r in range(sheet.nrows):
            values = [normalize_text(sheet.cell_value(r, c)) for c in range(sheet.ncols)]
            values = [v for v in values if v]
            if values:
                chunks.append(" | ".join(values))
    return "\n".join(chunks)


def _zip_child_is_skip(name: str, file_type: str) -> bool:
    lower = (name or "").lower()
    if any(lower.endswith(ext) for ext in ZIP_SKIP_NAME_EXT):
        return True
    if file_type in {"image", "rar", "7z"}:
        return True
    return False


def parse_zip(path: Path, extract_dir: Path) -> Dict[str, Any]:
    """ZIP：能解析出的文本都要；有足够文本即 text_read_ok。

    图纸/图片不挡成功；仅「文本类」子文件全失败且总文本不足时才失败。
    （修复 P1-02：原逻辑要求所有子文件都 read_ok，一张 dwg 就整包失败）
    """
    text_parts, children, total_size = [], [], 0
    extract_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as archive:
        infos = [item for item in archive.infolist() if not item.is_dir()]
        if len(infos) > MAX_ZIP_FILES:
            return {"text_read_ok": False, "text": "", "text_length": 0, "error": "ZIP文件数量超过限制"}
        for info in infos:
            total_size += info.file_size
            if total_size > MAX_ZIP_UNCOMPRESSED_BYTES:
                return {"text_read_ok": False, "text": "", "text_length": 0, "error": "ZIP解压总大小超过限制"}
            resolved = (extract_dir / info.filename).resolve()
            try:
                resolved.relative_to(extract_dir.resolve())
            except ValueError:
                children.append({
                    "name": info.filename, "text_read_ok": False,
                    "error": "ZIP路径穿越被拒绝", "skipped": True,
                })
                continue
            child_type = infer_file_type(info.filename)
            if _zip_child_is_skip(info.filename, child_type):
                children.append({
                    "name": info.filename, "type": child_type,
                    "text_read_ok": False, "text": "", "text_length": 0,
                    "error": "非文本类子文件，跳过", "skipped": True,
                })
                continue
            resolved.parent.mkdir(parents=True, exist_ok=True)
            with archive.open(info) as source:
                resolved.write_bytes(source.read())
            child = parse_attachment(resolved, child_type, allow_zip=False)
            child_rec = {"name": info.filename, "type": child_type, **child, "skipped": False}
            children.append(child_rec)
            if child.get("text_read_ok"):
                text_parts.append(f"[{info.filename}]\n{child.get('text', '')}")
    text = "\n\n".join(text_parts)
    meaningful = has_meaningful_text(text, MIN_ATTACHMENT_CHARS)
    text_unresolved = [
        item for item in children
        if not item.get("skipped")
        and not item.get("text_read_ok")
        and (item.get("type") in TEXT_BEARING_TYPES or infer_file_type(item.get("name", "")) in TEXT_BEARING_TYPES)
    ]
    # 有足够合并文本即成功；无文本时再报未读文本类文件数
    if meaningful:
        err = ""
        if text_unresolved:
            err = f"ZIP已提取正文；另有{len(text_unresolved)}个文本类子文件未读（不挡成功）"
        return {
            "text_read_ok": True,
            "text": text,
            "text_length": len(compact_text(text)),
            "error": err,
            "children": children,
            "ocr_required": any(item.get("ocr_required") for item in children if not item.get("skipped")),
        }
    return {
        "text_read_ok": False,
        "text": text,
        "text_length": len(compact_text(text)),
        "error": (
            f"ZIP内可解析文本不足"
            + (f"；文本类未读{len(text_unresolved)}个" if text_unresolved else "")
        ),
        "children": children,
        "ocr_required": any(item.get("ocr_required") for item in children if not item.get("skipped")),
    }


def attachment_is_optional_non_body(name: str, file_type: str) -> bool:
    """签章/委托/图片等：下载失败可记，文本失败不单独否定整项目 attachment_read_ok。"""
    n = name or ""
    if file_type in {"image"}:
        return True
    if any(h in n for h in OPTIONAL_ATTACH_NAME_HINTS):
        return True
    return False


def parse_attachment(path: Path, file_type: str, allow_zip: bool = True) -> Dict[str, Any]:
    try:
        if file_type == "docx":
            text = parse_docx(path)
        elif file_type == "doc":
            text = parse_doc(path)
        elif file_type == "pdf":
            text = parse_pdf(path)
        elif file_type == "xlsx":
            text = parse_xlsx(path)
        elif file_type == "xls":
            text = parse_xls(path)
        elif file_type == "zip" and allow_zip:
            return parse_zip(path, path.parent / f"{path.stem}_extracted")
        else:
            return {"text_read_ok": False, "text": "", "text_length": 0,
                    "error": f"暂不支持解析: {file_type}",
                    "ocr_required": file_type in {"pdf", "image"}}
        meaningful = has_meaningful_text(text, MIN_ATTACHMENT_CHARS)
        return {"text_read_ok": meaningful, "text": normalize_text(text),
                "text_length": len(compact_text(text)),
                "error": "" if meaningful else "附件有效文字不足，可能是扫描文件",
                "ocr_required": file_type in {"pdf", "image"} and not meaningful}
    except Exception as exc:
        return {"text_read_ok": False, "text": "", "text_length": 0,
                "error": f"附件解析失败: {exc}",
                "ocr_required": file_type in {"pdf", "image"}}


def _parse_ocr_progress_line(line: str) -> dict | None:
    """Parse tpk-ocr progress lines into {page, total, status} if possible."""
    text = (line or "").strip()
    if not text:
        return None
    # Preferred: [tpk-ocr] PROGRESS page=3 total=25 status=ocr_running
    m = re.search(
        r"PROGRESS\s+page=(\d+)\s+total=(\d+)\s+status=(\w+)",
        text,
        re.I,
    )
    if m:
        return {"page": int(m.group(1)), "total": int(m.group(2)), "status": m.group(3)}
    # Fallback: [tpk-ocr] OCR page 12...
    m2 = re.search(r"OCR page\s+(\d+)", text, re.I)
    if m2:
        return {"page": int(m2.group(1)), "total": 0, "status": "ocr_running"}
    if "Wrote:" in text or text.startswith("====="):
        return {"page": 0, "total": 0, "status": "writing"}
    return None


def _run_ocr_subprocess_with_progress(cmd: list[str], page_spec: str) -> Dict[str, Any]:
    """Run OCR with progress-aware timeout (not a single dead timeout).

    Policy (2026-07-14):
    - Poll ~every OCR_POLL_INTERVAL_SEC while reading stderr for page progress.
    - If no page progress for OCR_STALL_TIMEOUT_SEC (base 300s) → kill (stalled).
    - Absolute hard cap OCR_HARD_CAP_SEC (30 min) → kill.
    - If pages keep advancing, allow exceeding base until hard cap or completion.
    """
    start = time.time()
    last_progress_at = start
    last_page = 0
    last_status = "starting"
    stderr_chunks: list[str] = []
    stdout_chunks: list[str] = []

    print(
        f"  [OCR] start progress-watch stall={OCR_STALL_TIMEOUT_SEC}s "
        f"hard_cap={OCR_HARD_CAP_SEC}s pages={page_spec}",
        flush=True,
    )
    proc = subprocess.Popen(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        encoding="utf-8",
        errors="replace",
        bufsize=1,
    )
    assert proc.stderr is not None and proc.stdout is not None

    # Non-blocking-ish reads on Windows: use threads collecting pipes
    import threading
    import queue

    q_err: queue.Queue = queue.Queue()
    q_out: queue.Queue = queue.Queue()

    def _reader(stream, q: queue.Queue):
        try:
            for line in stream:
                q.put(line)
        except Exception:
            pass
        finally:
            q.put(None)

    t_err = threading.Thread(target=_reader, args=(proc.stderr, q_err), daemon=True)
    t_out = threading.Thread(target=_reader, args=(proc.stdout, q_out), daemon=True)
    t_err.start()
    t_out.start()

    err_done = out_done = False
    last_heartbeat = 0.0
    kill_reason = ""

    while True:
        now = time.time()
        elapsed = now - start

        # Drain queues
        while True:
            try:
                line = q_err.get_nowait()
            except queue.Empty:
                break
            if line is None:
                err_done = True
            else:
                stderr_chunks.append(line)
                # 任意 stderr 行都算「有活动」（单图 OCR 无页码时防误杀）
                last_progress_at = now
                prog = _parse_ocr_progress_line(line)
                if prog and prog.get("page", 0) > last_page:
                    last_page = int(prog["page"])
                    last_status = str(prog.get("status") or "ocr_running")
                    print(
                        f"  [OCR] progress page={last_page}"
                        f"{'/' + str(prog['total']) if prog.get('total') else ''} "
                        f"status={last_status} elapsed={int(elapsed)}s",
                        flush=True,
                    )
                elif prog and prog.get("status"):
                    last_status = str(prog["status"])

        while True:
            try:
                line = q_out.get_nowait()
            except queue.Empty:
                break
            if line is None:
                out_done = True
            else:
                stdout_chunks.append(line)

        rc = proc.poll()
        if rc is not None and err_done and out_done:
            break
        if rc is not None and (err_done or out_done):
            # give a short moment for the other stream
            if now - start > elapsed + 2:
                break

        # Heartbeat every ~30s even without new pages
        if now - last_heartbeat >= OCR_POLL_INTERVAL_SEC:
            stall_for = int(now - last_progress_at)
            print(
                f"  [OCR] heartbeat elapsed={int(elapsed)}s page={last_page} "
                f"status={last_status} stall_for={stall_for}s",
                flush=True,
            )
            last_heartbeat = now

        if elapsed >= OCR_HARD_CAP_SEC:
            kill_reason = f"hard_cap ({OCR_HARD_CAP_SEC}s)"
            break
        # Base / stall: 连续无进展超过 STALL 则停（有进展则可超过 base 直到 hard_cap）
        stalled = (now - last_progress_at) >= OCR_STALL_TIMEOUT_SEC
        if stalled and page_spec == "image":
            # 单图无页码：以 stderr 活动为进展；完全静默 STALL 则停
            kill_reason = f"image_ocr_stall ({OCR_STALL_TIMEOUT_SEC}s no activity)"
            break
        if stalled and last_page == 0 and elapsed >= OCR_STALL_TIMEOUT_SEC:
            kill_reason = f"no_progress_base ({OCR_STALL_TIMEOUT_SEC}s, never advanced page)"
            break
        if stalled and last_page > 0:
            kill_reason = (
                f"stalled after page={last_page} "
                f"(no advance for {OCR_STALL_TIMEOUT_SEC}s)"
            )
            break

        time.sleep(0.5)

    if kill_reason and proc.poll() is None:
        print(f"  [OCR] terminating: {kill_reason}", flush=True)
        try:
            proc.kill()
        except Exception:
            pass
        try:
            proc.wait(timeout=10)
        except Exception:
            pass
        return {
            "text_read_ok": False,
            "text": "",
            "text_length": 0,
            "error": f"OCR stopped: {kill_reason}; last_page={last_page}; pages={page_spec}",
            "ocr_executed": True,
            "ocr_pages": page_spec,
            "ocr_last_page": last_page,
            "ocr_elapsed_sec": int(time.time() - start),
            "ocr_status": "timeout_or_stall",
        }

    # Process finished normally — join readers
    t_err.join(timeout=5)
    t_out.join(timeout=5)
    while not err_done:
        try:
            line = q_err.get(timeout=0.2)
        except queue.Empty:
            break
        if line is None:
            err_done = True
        else:
            stderr_chunks.append(line)
    while not out_done:
        try:
            line = q_out.get(timeout=0.2)
        except queue.Empty:
            break
        if line is None:
            out_done = True
        else:
            stdout_chunks.append(line)

    rc = proc.returncode if proc.returncode is not None else -1
    stderr_text = "".join(stderr_chunks)
    ocr_text = "".join(stdout_chunks).strip()
    elapsed = int(time.time() - start)

    if rc != 0:
        return {
            "text_read_ok": False,
            "text": "",
            "text_length": 0,
            "error": f"OCR failed rc={rc}: {(stderr_text or '')[:200]}",
            "ocr_executed": True,
            "ocr_pages": page_spec,
            "ocr_last_page": last_page,
            "ocr_elapsed_sec": elapsed,
            "ocr_status": "failed",
        }

    if not ocr_text:
        return {
            "text_read_ok": False,
            "text": "",
            "text_length": 0,
            "error": "OCR result empty",
            "ocr_executed": True,
            "ocr_pages": page_spec,
            "ocr_last_page": last_page,
            "ocr_elapsed_sec": elapsed,
            "ocr_status": "empty",
        }

    if ocr_text.startswith("# OCR Result"):
        lines = ocr_text.split("\n")
        text_start = 0
        for i, line in enumerate(lines):
            if line.startswith("## Text"):
                text_start = i + 2
                break
        if text_start > 0:
            ocr_text = "\n".join(lines[text_start:]).strip()

    meaningful = has_meaningful_text(ocr_text, MIN_ATTACHMENT_CHARS)
    return {
        "text_read_ok": meaningful,
        "text": normalize_text(ocr_text) if meaningful else "",
        "text_length": len(compact_text(ocr_text)),
        "error": "" if meaningful else "OCR result has insufficient text",
        "ocr_executed": True,
        "ocr_pages": page_spec,
        "ocr_last_page": last_page,
        "ocr_elapsed_sec": elapsed,
        "ocr_status": "ok" if meaningful else "weak_text",
    }


def run_ocr_on_attachment(path: Path, file_type: str) -> Dict[str, Any]:
    """Call tpk-ocr skill to OCR scanned PDF or image attachment.
    
    Returns dict with text_read_ok, text, text_length, ocr_executed, error.
    2026-07-14: progress heartbeat + stall base 300s + hard cap 30min; skip only >100MB.
    """
    # Locate tpk-ocr scripts (runtime: 03_脚本工具/tpk-ocr/scripts; skill: scripts/tpk-ocr/scripts)
    here = Path(__file__).resolve().parent
    candidates = [
        here / "tpk-ocr" / "scripts",
        here.parent / "scripts" / "tpk-ocr" / "scripts",
        Path.home() / ".qclaw" / "skills" / "tpk-ocr" / "scripts",
        Path.home() / ".agents" / "skills" / "tpk-ocr" / "scripts",
        Path.home() / ".grok" / "skills" / "tpk-ocr" / "scripts",
    ]
    skill_dir = next((p for p in candidates if p.is_dir()), None)
    if skill_dir is None:
        return {"text_read_ok": False, "text": "", "text_length": 0,
                "error": "tpk-ocr skill not found (tried runtime + user skill paths)",
                "ocr_executed": False}
    
    try:
        file_size = path.stat().st_size if path.exists() else 0
    except Exception:
        file_size = 0

    if file_size > LARGE_FILE_SKIP_OCR:
        return {
            "text_read_ok": False,
            "text": "",
            "text_length": 0,
            "error": (
                f"file too large for OCR ({file_size // 1024 // 1024}MB "
                f"> {LARGE_FILE_SKIP_OCR // 1024 // 1024}MB skip threshold)"
            ),
            "ocr_executed": False,
            "ocr_status": "skipped_large",
        }

    # 页数：大文件仍只取前面若干页（够业务定性）；超时改由进度心跳控制
    if file_size > 40 * 1024 * 1024:
        page_spec = "1-15"
    elif file_size > 10 * 1024 * 1024:
        page_spec = "1-25"
    else:
        page_spec = "1-40"

    if file_type == "pdf":
        ocr_script = skill_dir / "ocr_pdf.py"
        cmd = [
            sys.executable, str(ocr_script), str(path),
            "--pages", page_spec, "--dpi", "200", "--lang", "ch", "--format", "text",
        ]
    else:
        ocr_script = skill_dir / "ocr_image.py"
        cmd = [sys.executable, str(ocr_script), str(path), "--lang", "ch", "--format", "text"]
        page_spec = "image"
    
    if not ocr_script.exists():
        return {"text_read_ok": False, "text": "", "text_length": 0,
                "error": f"OCR script not found: {ocr_script}", "ocr_executed": False}
    
    try:
        return _run_ocr_subprocess_with_progress(cmd, page_spec)
    except Exception as exc:
        return {"text_read_ok": False, "text": "", "text_length": 0,
                "error": f"OCR exception: {exc}", "ocr_executed": True}


def attachment_is_required(project: Dict[str, Any], body: Dict[str, Any],
                           attachments: List[Dict[str, Any]]) -> bool:
    return bool(project.get("attachment_required")) or bool(
        attachments and (not body.get("body_read_ok") or
                         any(hint in body.get("body_text", "") for hint in ATTACHMENT_HINTS))
    )


def empty_results(project: Dict[str, Any], status: str,
                  reason: str) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    bid_id, title, url = extract_bid_id(project), normalize_text(project.get("title")), project_url(project)
    vip = {"bid_id": bid_id, "title": title, "url": url, "status": status,
           "login_ok": status not in {"need_login", "captcha"}, "body_read_ok": False,
           "body_identity_ok": False, "body_text": "", "body_text_length": 0,
           "body_source": "", "body_invalid_reason": reason, "error": reason}
    attachment = {"bid_id": bid_id, "title": title, "url": url,
                  "attachment_available": False,
                  "attachment_required": bool(project.get("attachment_required")),
                  "attachment_read_ok": False, "attachment_status": status,
                  "attachment_count": 0, "attachment_read_ok_count": 0,
                  "attachment_unresolved_count": 0, "attachments": [], "error": reason}
    return vip, attachment


def process_project(client: CDPClient, project: Dict[str, Any],
                    download_root: Path) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    title, url, bid_id = normalize_text(project.get("title")), project_url(project), extract_bid_id(project)
    proj_start = time.time()
    slog(f"[PROJECT] begin bid={bid_id or '-'} title={title[:60]}")
    if not url:
        return empty_results(project, "missing_detail_url", "项目缺少详情页URL")
    try:
        client.navigate(url)
    except Exception as exc:
        return empty_results(project, "error", str(exc))
    gate, gate_reason = page_gate(client)
    if gate != "ok":
        client.screenshot(f"{gate}_{bid_id or 'unknown'}")
        slog(f"[PROJECT] gate={gate} reason={gate_reason}")
        return empty_results(project, gate, gate_reason)

    body = extract_body(client, title)
    vip = {"bid_id": bid_id, "title": title, "url": url,
           "status": "ok" if body["body_read_ok"] else "body_empty",
           "login_ok": True, **body,
           "error": "" if body["body_read_ok"] else body["body_invalid_reason"]}
    slog(f"[PROJECT] body_ok={body.get('body_read_ok')} status={vip.get('status')}")

    timeline = extract_progress_timeline(client, title)
    vip["progress_timeline"] = timeline

    detected = detect_attachments(client)
    required = attachment_is_required(project, body, detected)
    project_dir = download_root / safe_filename(bid_id or title or "unknown")
    attachment_items = []
    project_deadline = proj_start + PER_PROJECT_BUDGET_SEC
    slog(
        f"[PROJECT] attachments detected={len(detected)} required={required} "
        f"project_budget={PER_PROJECT_BUDGET_SEC}s attach_budget={PER_ATTACHMENT_BUDGET_SEC}s"
    )
    for index, item in enumerate(detected, start=1):
        if time.time() >= project_deadline:
            slog(f"[PROJECT] project budget exhausted before attach #{index}")
            attachment_items.append({
                "name": normalize_text(item.get("name")) or f"附件{index}",
                "type": "unknown",
                "download_ok": False,
                "download_path": "",
                "source_url": item.get("href", ""),
                "real_url": "",
                "content_type": "",
                "sha256": "",
                "text_read_ok": False,
                "text": "",
                "text_length": 0,
                "error": f"project budget exhausted ({PER_PROJECT_BUDGET_SEC}s)",
                "ocr_required": False,
            })
            continue
        attach_deadline = min(project_deadline, time.time() + PER_ATTACHMENT_BUDGET_SEC)
        t_att = time.time()
        downloaded = download_attachment(
            client, item, project_dir, index, url, budget_deadline=attach_deadline
        )
        if time.time() >= attach_deadline and not downloaded.get("download_ok"):
            downloaded = {
                "download_ok": False,
                "error": downloaded.get("error")
                or f"attachment budget exhausted ({PER_ATTACHMENT_BUDGET_SEC}s)",
            }
        file_type = downloaded.get("file_type") or infer_file_type(
            item.get("name", ""), downloaded.get("content_type", ""))
        parsed = (parse_attachment(Path(downloaded["download_path"]), file_type)
                  if downloaded.get("download_ok") else
                  {"text_read_ok": False, "text": "", "text_length": 0,
                   "error": downloaded.get("error", "下载失败"),
                   "ocr_required": file_type in {"pdf", "image"}})
        slog(
            f"  [ATTACH] #{index} download_ok={downloaded.get('download_ok')} "
            f"type={file_type} parse_ok={parsed.get('text_read_ok')} "
            f"dt={time.time()-t_att:.1f}s"
        )
        # OCR fallback: if attachment needs OCR and text extraction failed, run tpk-ocr
        if (parsed.get("ocr_required") and not parsed.get("text_read_ok")
                and downloaded.get("download_ok") and file_type in {"pdf", "image"}):
            if time.time() >= project_deadline or time.time() >= attach_deadline:
                parsed["error"] = "OCR skipped: budget exhausted"
                slog(f"  [OCR] skip budget exhausted name={item.get('name', '')[:60]}")
            else:
                slog(f"  [OCR] 扫描件附件，调用 tpk-ocr: {item.get('name', '')[:60]}")
                ocr_result = run_ocr_on_attachment(Path(downloaded["download_path"]), file_type)
                parsed["ocr_executed"] = ocr_result.get("ocr_executed", False)
                parsed["ocr_elapsed_sec"] = ocr_result.get("ocr_elapsed_sec")
                parsed["ocr_status"] = ocr_result.get("ocr_status")
                if ocr_result.get("text_read_ok"):
                    parsed["text_read_ok"] = True
                    parsed["text"] = ocr_result["text"]
                    parsed["text_length"] = ocr_result["text_length"]
                    parsed["ocr_required"] = False
                    parsed["error"] = ""
                else:
                    parsed["error"] = ocr_result.get("error", "OCR failed")
                slog(
                    f"  [OCR] done ok={ocr_result.get('text_read_ok')} "
                    f"status={ocr_result.get('ocr_status')} "
                    f"elapsed={ocr_result.get('ocr_elapsed_sec')}s"
                )
        name_disp = normalize_text(item.get("name")) or f"附件{index}"
        optional = attachment_is_optional_non_body(name_disp, file_type)
        attachment_items.append({
            "name": name_disp,
            "type": file_type, "download_ok": bool(downloaded.get("download_ok")),
            "download_path": downloaded.get("download_path", ""),
            "source_url": downloaded.get("source_url", item.get("href", "")),
            "real_url": downloaded.get("real_url", ""),
            "content_type": downloaded.get("content_type", ""),
            "sha256": downloaded.get("sha256", ""),
            "optional_non_body": optional,
            **parsed,
        })
    successful = [item for item in attachment_items if item.get("text_read_ok")]
    # 阻塞性未读：下载失败的文本主件，或解析失败且非签章/图片类
    blocking_unresolved = []
    soft_unresolved = []
    for item in attachment_items:
        if item.get("text_read_ok"):
            continue
        name = item.get("name") or ""
        ft = item.get("type") or ""
        optional = bool(item.get("optional_non_body")) or attachment_is_optional_non_body(name, ft)
        if not item.get("download_ok"):
            # 必下失败：主文本类算阻塞；可选类算 soft
            if optional:
                soft_unresolved.append(item)
            else:
                blocking_unresolved.append(item)
            continue
        if optional:
            soft_unresolved.append(item)
            continue
        if ft in TEXT_BEARING_TYPES or ft == "zip":
            blocking_unresolved.append(item)
        else:
            soft_unresolved.append(item)
    # complete：至少有一份正文级文本，且无阻塞性缺口
    complete = bool(successful) and not blocking_unresolved
    # 仅有可选附件失败、或主件已读：complete
    status = (
        "no_attachment" if not detected else
        "complete" if complete else
        "partial" if successful else
        "unreadable"
    )
    attachment = {
        "bid_id": bid_id, "title": title, "url": url,
        "attachment_available": bool(detected), "attachment_required": required,
        "attachment_read_ok": complete, "attachment_status": status,
        "attachment_count": len(attachment_items),
        "attachment_read_ok_count": len(successful),
        "attachment_unresolved_count": len(blocking_unresolved) + len(soft_unresolved),
        "attachment_blocking_unresolved_count": len(blocking_unresolved),
        "attachment_soft_unresolved_count": len(soft_unresolved),
        "attachments": attachment_items,
        "error": (
            "" if complete or not required
            else f"有{len(blocking_unresolved)}个主附件未形成可用证据"
            + (f"（另{len(soft_unresolved)}个签章/可选附件可忽略）" if soft_unresolved else "")
        ),
        "project_elapsed_sec": round(time.time() - proj_start, 1),
    }
    slog(
        f"[PROJECT] end bid={bid_id or '-'} attach_status={status} "
        f"ok={len(successful)}/{len(attachment_items)} "
        f"elapsed={attachment['project_elapsed_sec']}s"
    )
    client.screenshot(f"project_{bid_id or hashlib.md5(title.encode('utf-8')).hexdigest()[:8]}")
    return vip, attachment


def write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Qianlima CDP正文和附件读取器 v05")
    parser.add_argument("--input", required=True, help="Phase 2A筛选JSON")
    parser.add_argument("--output-dir", required=True, help="输出目录")
    parser.add_argument("--date", required=True, help="业务日期 YYYYMMDD")
    # Chrome 150+ DevTools HTTP only accepts Host: localhost (127.0.0.1 returns 404).
    parser.add_argument("--cdp", default="http://localhost:9222", help="Chrome CDP地址（须用 localhost）")
    parser.add_argument("--max-projects", type=int, default=0)
    parser.add_argument("--config", help="本机敏感配置；仅用于登录恢复，不写入输出")
    parser.add_argument(
        "--login-script",
        default=str(Path(__file__).with_name("qianlima_auto_login_20260622_v02.py")),
        help="自动登录脚本路径",
    )
    args = parser.parse_args()
    global _SEG_LOG
    input_path, output_dir = Path(args.input), Path(args.output_dir)
    screenshot_dir, download_dir = output_dir / "screenshots", output_dir / "downloads"
    screenshot_dir.mkdir(parents=True, exist_ok=True)
    download_dir.mkdir(parents=True, exist_ok=True)
    # P0：分段日志 → runs/<run_id>/99_logs/phase2c_attach_*.log
    log_dir = resolve_phase2_log_dir(output_dir)
    log_path = log_dir / f"phase2c_attach_{args.date}_{time.strftime('%H%M%S')}.log"
    _SEG_LOG = SegmentLogger(log_path)
    slog(f"segment_log={log_path}")
    slog(f"cdp={args.cdp} input={input_path} output_dir={output_dir}")

    projects = iter_projects(json.loads(input_path.read_text(encoding="utf-8-sig")))
    if args.max_projects > 0:
        projects = projects[:args.max_projects]

    vip_path = output_dir / f"VIP原文阅读_CDP_{args.date}_v05.json"
    attachment_path = output_dir / f"附件预览正文读取_CDP_{args.date}_v05.json"
    summary_path = output_dir / f"CDP正文附件读取摘要_{args.date}_v05.json"
    vip_results, attachment_results = [], []
    legal_empty = False

    try:
        if not projects:
            # Legal full-exclude: Phase 2A left no recommended/considered/low_score rows.
            slog("No projects selected from input; writing legal empty VIP/attachment JSON.")
            legal_empty = True
        else:
            client = CDPClient(args.cdp, screenshot_dir)
            try:
                client.connect()
                login_script = Path(args.login_script)
                config_path = Path(args.config) if args.config else None
                for index, project in enumerate(projects, start=1):
                    title = normalize_text(project.get("title"))
                    slog(f"[{index}/{len(projects)}] {title[:60]}")
                    try:
                        vip, attachment = process_project(client, project, download_dir)
                        initial_status = vip.get("status")
                        recovery_attempted = False

                        if initial_status == "need_login":
                            recovery_attempted = True
                            slog(f"[{index}/{len(projects)}] need_login -> restore once")
                            restored = restore_login_once(client, login_script, config_path)
                            if restored:
                                vip, attachment = process_project(client, project, download_dir)
                            else:
                                client.screenshot(f"need_login_restore_failed_{extract_bid_id(project) or index}")
                        elif initial_status == "body_empty":
                            recovery_attempted = True
                            slog(f"[{index}/{len(projects)}] body_empty -> reload once")
                            vip, attachment = process_project(client, project, download_dir)

                        vip["recovery_attempted"] = recovery_attempted
                        vip["initial_status"] = initial_status
                        if recovery_attempted and vip.get("status") in {"need_login", "body_empty"}:
                            client.screenshot(
                                f"{vip.get('status')}_after_retry_{extract_bid_id(project) or index}"
                            )
                    except Exception as exc:
                        traceback.print_exc()
                        slog(f"[{index}/{len(projects)}] UNCAUGHT {exc}")
                        vip, attachment = empty_results(project, "error", f"未捕获异常: {exc}")
                    vip_results.append(vip)
                    attachment_results.append(attachment)
            finally:
                client.close()
    finally:
        if _SEG_LOG is not None:
            _SEG_LOG.close()
            _SEG_LOG = None

    summary = {
        "ok": (legal_empty or len(vip_results) == len(projects) == len(attachment_results)),
        "legal_empty": legal_empty,
        "input_count": len(projects), "vip_output_count": len(vip_results),
        "attachment_output_count": len(attachment_results),
        "body_read_ok_count": sum(bool(item.get("body_read_ok")) for item in vip_results),
        "attachment_required_count": sum(bool(item.get("attachment_required")) for item in attachment_results),
        "attachment_read_ok_count": sum(bool(item.get("attachment_read_ok")) for item in attachment_results),
        "need_login_count": sum(item.get("status") == "need_login" for item in vip_results),
        "captcha_count": sum(item.get("status") == "captcha" for item in vip_results),
        "vip_json": str(vip_path), "attachment_json": str(attachment_path),
    }
    vip_projects = []
    for item, attachment in zip(vip_results, attachment_results):
        vip_projects.append({
            **item,
            "content": item.get("body_text", ""),
            "content_preview": item.get("body_text", "")[:500],
            "body_clean": item.get("body_text", ""),
            "text_length": item.get("body_text_length", 0),
            "content_source": item.get("body_source", ""),
            "attachments": [
                {"name": child.get("name", ""), "type": child.get("type", "")}
                for child in attachment.get("attachments", [])
            ],
            "attachment_preview_required": attachment.get("attachment_required", False),
            "attachment_reason": attachment.get("error", ""),
        })
    attachment_projects = []
    for item in attachment_results:
        attachment_projects.append({
            **item,
            "target": {
                "bid_id": item.get("bid_id", ""),
                "title": item.get("title", ""),
                "url": item.get("url", ""),
            },
            "status": "ok" if item.get("attachment_read_ok") else item.get("attachment_status", "unreadable"),
            "attachment_results": [
                {
                    **child,
                    "capture": {
                        "text": child.get("text", ""),
                        "textPreview": child.get("text", "")[:500],
                    },
                    "pdf_ocr": {
                        "enabled": False,
                        "required": bool(child.get("ocr_required")),
                        "text": "",
                    },
                }
                for child in item.get("attachments", [])
            ],
        })
    vip_payload = {
        "version": "CDP-v05",
        "generated_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
        "source": "Qianlima CDP Playwright",
        "total": len(vip_projects),
        "legal_empty": legal_empty,
        "summary": {
            "ok": summary["body_read_ok_count"],
            "need_login": summary["need_login_count"],
            "captcha": summary["captcha_count"],
            "error": sum(item.get("status") == "error" for item in vip_projects),
            "body_empty": sum(item.get("status") == "body_empty" for item in vip_projects),
        },
        "projects": vip_projects,
    }
    attachment_payload = {
        "version": "CDP-v05",
        "generated_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
        "source_vip_json": str(vip_path),
        "targets_total": len(attachment_projects),
        "legal_empty": legal_empty,
        "summary": {
            "ok": summary["attachment_read_ok_count"],
            "partial": sum(item.get("attachment_status") == "partial" for item in attachment_projects),
            "unreadable": sum(item.get("attachment_status") == "unreadable" for item in attachment_projects),
            "no_attachment": sum(item.get("attachment_status") == "no_attachment" for item in attachment_projects),
        },
        "results": attachment_projects,
    }
    write_json(vip_path, vip_payload)
    write_json(attachment_path, attachment_payload)
    write_json(summary_path, summary)
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    # Count-matched output (or legal empty) is success; per-item gaps go to validator/rejudge.
    return 0 if summary["ok"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
